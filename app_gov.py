import streamlit as st
import os
import pandas as pd
from docx import Document

# --- CONFIGURAÇÃO DE CAMINHOS ---
CAMINHO_BASE = r"C:\LicitFlow_GOV\Processos_Licitatorios"
if not os.path.exists(CAMINHO_BASE):
    os.makedirs(CAMINHO_BASE)

# --- MOTOR DE IA (MINUTAS) ---
def gerador_ia_completo(tipo_doc, nome_obra, descricao):
    # Dicionário de conteúdos pré-estruturados por tipo de documento
    if tipo_doc == "PLANO_TRABALHO":
        corpo = f"""1. OBJETO E JUSTIFICATIVA:
A presente proposta visa a execução de {nome_obra}, buscando solucionar as patologias e necessidades identificadas no setor.

2. METAS E ETAPAS:
- META 1: Mobilização e Instalação de Canteiro (Previsão: 10 dias).
- META 2: Execução de Infraestrutura e Serviços Preliminares.
- META 3: Execução de Alvenaria, Acabamentos e Instalações.
- META 4: Limpeza final e entrega do objeto.

3. CRONOGRAMA DE DESEMBOLSO:
Os pagamentos serão efetuados mediante a aferição das medições mensais, seguindo o cronograma físico-financeiro aprovado pela fiscalização, garantindo o equilíbrio financeiro da execução.

4. RESULTADOS ESPERADOS:
Entrega da unidade plenamente funcional, atendendo às normas de acessibilidade e segurança vigentes."""
    
    elif tipo_doc == "PROJETO_BASICO":
        corpo = f"""1. ELEMENTOS DO PROJETO:
O Projeto Básico para {nome_obra} contempla os elementos necessários e suficientes, com nível de precisão adequado, para caracterizar a obra.

2. LEVANTAMENTOS:
Realização de vistoria técnica in loco para conferência de medidas e quantitativos.

3. ESPECIFICAÇÕES:
Os materiais deverão seguir o padrão de qualidade exigido pela Secretaria, com foco na durabilidade e baixa manutenção."""

    else:
        corpo = f"Diretrizes para {nome_obra}. Foco em normas técnicas e legislação 14.133."

    minuta = f"""[MINUTA TÉCNICA - LICITFLOW GOV]
OBRA: {nome_obra}
DOCUMENTO: {tipo_doc.replace('_', ' ')}

{corpo}

--------------------------------------------------
ESTA MINUTA DEVE SER CONFERIDA E EDITADA PELA ENGENHARIA.
"""
    return minuta

def criar_estrutura_obra(nome_obra):
    nome_limpo = nome_obra.replace(" ", "_").upper()
    caminho_obra = os.path.join(CAMINHO_BASE, nome_limpo)
    subpastas = ["01_Planejamento", "02_Projetos", "03_Orcamento", "04_Fiscalizacao"]
    if not os.path.exists(caminho_obra):
        os.makedirs(caminho_obra)
        for sub in subpastas:
            os.makedirs(os.path.join(caminho_obra, sub))
    return caminho_obra

def main():
    st.set_page_config(page_title="LicitFlow Gov AI", layout="wide")
    st.sidebar.title("🏛️ LicitFlow Gov AI")
    
    # SELEÇÃO DE OBRA
    st.sidebar.subheader("Seleção de Obra")
    obras = [f for f in os.listdir(CAMINHO_BASE) if os.path.isdir(os.path.join(CAMINHO_BASE, f))]
    if obras:
        sel = st.sidebar.selectbox("Trabalhar na Obra:", ["-- Selecione --"] + obras)
        if sel != "-- Selecione --":
            st.session_state['pasta_ativa'] = os.path.join(CAMINHO_BASE, sel)
    
    menu = st.sidebar.radio("Navegação:", ["Nova Demanda (DFD/ETP)", "Projetos e Planos", "Orçamento/Cronograma", "Acompanhamento"])

    # --- MÓDULO 1: FASE PREPARATÓRIA (DFD, ETP E TERMO DE REFERÊNCIA) ---
    if menu == "Nova Demanda (DFD/ETP)":
        st.header("📘 Planejamento e Instrução Processual")
        
        with st.form("form_fase_preparatoria"):
            nome_obra = st.text_input("Nome da Obra/Serviço (Objeto):", placeholder="Ex: Reforma da Unidade de Saúde Central")
            problema = st.text_area("Justificativa / Descrição da Necessidade:", placeholder="Descreva os problemas detectados e o que precisa ser feito...")
            
            st.write("---")
            st.subheader("🤖 Geração de Minutas via IA")
            
            col_ia1, col_ia2, col_ia3 = st.columns(3)
            
            with col_ia1:
                if st.form_submit_button("📄 Gerar DFD"):
                    st.session_state['dfd_ia'] = f"DOCUMENTO DE FORMALIZAÇÃO DE DEMANDA\n\n1. NECESSIDADE: {problema}\n2. ALINHAMENTO: Previsto no PCA.\n3. RESPONSÁVEL: Secretaria de Infraestrutura."
            
            with col_ia2:
                if st.form_submit_button("📑 Gerar ETP"):
                    st.session_state['etp_ia'] = gerador_ia_completo("ETP", nome_obra, problema)
            
            with col_ia3:
                if st.form_submit_button("📜 Gerar Termo de Ref. (TR)"):
                    st.session_state['tr_ia'] = f"""TERMO DE REFERÊNCIA - LEI 14.133/21
OBRA: {nome_obra}

1. OBJETO: Contratação de empresa de engenharia para {nome_obra}.
2. DESCRIÇÃO: Detalhamento conforme Projetos e Memorial Descritivo.
3. REQUISITOS: Qualificação técnica e operacional.
4. FISCALIZAÇÃO: Acompanhamento via Diário de Obra e Medições mensais.
5. PAGAMENTO: Conforme cronograma físico-financeiro aprovado."""

            st.write("---")
            # Campos de Edição
            txt_dfd = st.text_area("Edição DFD:", value=st.session_state.get('dfd_ia', ""), height=150)
            txt_etp = st.text_area("Edição ETP:", value=st.session_state.get('etp_ia', ""), height=150)
            txt_tr = st.text_area("Edição Termo de Referência:", value=st.session_state.get('tr_ia', ""), height=150)
            
            btn_salvar = st.form_submit_button("🔨 Consolidar e Salvar Fase Preparatória")
            
            if btn_salvar and nome_obra:
                p_obra = criar_estrutura_obra(nome_obra)
                st.session_state['pasta_ativa'] = p_obra
                p_plan = os.path.join(p_obra, "01_Planejamento")
                
                # DOC 1: DFD
                d_dfd = Document(); d_dfd.add_heading('DFD', 0)
                for p in txt_dfd.split('\n'): d_dfd.add_paragraph(p)
                d_dfd.save(os.path.join(p_plan, "01_DFD.docx"))
                
                # DOC 2: ETP
                d_etp = Document(); d_etp.add_heading('ETP', 0)
                for p in txt_etp.split('\n'): d_etp.add_paragraph(p)
                d_etp.save(os.path.join(p_plan, "02_ETP.docx"))
                
                # DOC 3: TR (A Peça que faltava)
                d_tr = Document(); d_tr.add_heading('TERMO DE REFERÊNCIA', 0)
                for p in txt_tr.split('\n'): d_tr.add_paragraph(p)
                d_tr.save(os.path.join(p_plan, "03_TERMO_DE_REFERENCIA.docx"))
                
                st.success(f"✅ Documentos salvos com sucesso na pasta: {p_plan}")
                st.rerun()

    # --- MÓDULO 2: PROJETOS, MEMORIAIS E PLANO DE TRABALHO ---
    elif menu == "Projetos e Planos":
        if 'pasta_ativa' not in st.session_state:
            st.warning("⚠️ Selecione uma obra na barra lateral antes de prosseguir.")
        else:
            pasta_ativa = st.session_state['pasta_ativa']
            nome_obra = os.path.basename(pasta_ativa)
            st.header(f"📂 Gestão Técnica: {nome_obra}")
            
            t1, t2, t3 = st.tabs(["🏗️ Engenharia (PB/PE/Memorial)", "📋 Plano de Trabalho", "📤 Repositório de Desenhos"])
            
            with t1:
                st.subheader("Documentos Técnicos de Engenharia")
                # Incluímos o Memorial Descritivo no seletor
                tipo_doc_tec = st.radio(
                    "Selecione o documento para gerar:", 
                    ["Projeto Básico", "Projeto Executivo", "Memorial Descritivo"],
                    horizontal=True
                )
                
                if st.button(f"🤖 Gerar Minuta de {tipo_doc_tec}"):
                    if tipo_doc_tec == "Memorial Descritivo":
                        st.session_state['minuta_tec'] = f"""MEMORIAL DESCRITIVO E ESPECIFICAÇÕES TÉCNICAS
OBRA: {nome_obra}

1. OBJETO: Descrição detalhada dos materiais e métodos para {nome_obra}.
2. PADRÕES DE QUALIDADE: Os materiais (cimentos, tintas, revestimentos) devem atender às NBRs e especificações da Secretaria.
3. EXECUÇÃO: Procedimentos para fundação, alvenaria, cobertura e acabamentos.
4. NORMAS DE SEGURANÇA: Obediência à NR-18 e uso obrigatório de EPIs.
5. LIMPEZA FINAL: A obra deve ser entregue livre de entulhos e com limpeza fina realizada."""
                    else:
                        # Chama a função de IA genérica para PB ou PE
                        cod_ia = "PROJETO_BASICO" if tipo_doc_tec == "Projeto Básico" else "PROJETO_EXECUTIVO"
                        st.session_state['minuta_tec'] = gerador_ia_completo(cod_ia, nome_obra, "Diretrizes Técnicas")
                
                # Campo de edição para o engenheiro complementar
                txt_tec_edit = st.text_area(
                    f"Conteúdo do {tipo_doc_tec}:", 
                    value=st.session_state.get('minuta_tec', ""), 
                    height=300
                )
                
                if st.button(f"💾 Salvar {tipo_doc_tec}"):
                    d = Document()
                    d.add_heading(f"{tipo_doc_tec.upper()} - {nome_obra}", 0)
                    for p in txt_tec_edit.split('\n'):
                        d.add_paragraph(p)
                    
                    # Nome do arquivo amigável
                    nome_arq = f"{tipo_doc_tec.replace(' ', '_').upper()}.docx"
                    path_final = os.path.join(pasta_ativa, "02_Projetos", nome_arq)
                    d.save(path_final)
                    st.success(f"✅ {tipo_doc_tec} salvo com sucesso em: {path_final}")

            with t2:
                st.subheader("Plano de Trabalho Detalhado")
                if st.button("🤖 Gerar Rascunho do Plano de Trabalho"):
                    st.session_state['plano_ia'] = gerador_ia_completo("PLANO_TRABALHO", nome_obra, "Metas e Cronograma")
                
                txt_plano = st.text_area("Edição do Plano:", value=st.session_state.get('plano_ia', ""), height=300)
                
                if st.button("💾 Salvar Plano de Trabalho"):
                    d = Document()
                    d.add_heading(f"PLANO DE TRABALHO - {nome_obra}", 0)
                    for p in txt_plano.split('\n'):
                        d.add_paragraph(p)
                    
                    path_pt = os.path.join(pasta_ativa, "02_Projetos", "PLANO_DE_TRABALHO.docx")
                    d.save(path_pt)
                    st.success(f"✅ Plano de Trabalho salvo em: {path_pt}")

            with t3:
                st.subheader("Upload de Pranchas e Memoriais Externos")
                up_arqs = st.file_uploader("Arraste os arquivos de projeto (PDF, DWG, XLSX):", accept_multiple_files=True)
                if st.button("🚀 Processar e Organizar"):
                    if up_arqs:
                        for a in up_arqs:
                            with open(os.path.join(pasta_ativa, "02_Projetos", a.name), "wb") as f:
                                f.write(a.getbuffer())
                        st.success("Documentos técnicos anexados à pasta de projetos!")

    # --- MÓDULO 3: ORÇAMENTO, TABELAS E CRONOGRAMA ---
    elif menu == "Orçamento/Cronograma":
        if 'pasta_ativa' not in st.session_state:
            st.warning("⚠️ Selecione uma obra na barra lateral para acessar o orçamento.")
        else:
            pasta_ativa = st.session_state['pasta_ativa']
            nome_obra = os.path.basename(pasta_ativa)
            st.header(f"💰 Engenharia de Custos: {nome_obra}")

            tab_base, tab_planilha, tab_cron = st.tabs([
                "📋 Tabelas de Referência (SINAPI/SICRO)", 
                "📝 Montagem da Planilha Orçamentária", 
                "📅 Cronograma Físico-Financeiro"
            ])

            # --- ABA 1: CONFIGURAÇÃO DA BASE ---
            with tab_base:
                st.subheader("Configuração da Base de Preços")
                col_b1, col_b2 = st.columns(2)
                with col_b1:
                    fonte_ref = st.selectbox("Fonte:", ["SINAPI", "SICRO", "CDHU", "FDE", "Tabela Própria"])
                    mes_ano = st.text_input("Mês/Ano de Referência:", placeholder="Ex: 01/2026")
                with col_b2:
                    encargos = st.radio("Encargos Sociais:", ["Desonerado", "Não Desonerado"], horizontal=True)
                    bdi_obra = st.number_input("BDI da Obra (%):", min_value=0.0, max_value=50.0, value=25.0)

                up_tabela = st.file_uploader(f"Upload da Tabela {fonte_ref} (.xlsx):", type=["xlsx"])
                if st.button("📥 Registrar Tabela no Processo"):
                    if up_tabela:
                        path_tabela = os.path.join(pasta_ativa, "03_Orcamento", f"BASE_REFERENCIA_{fonte_ref}.xlsx")
                        with open(path_tabela, "wb") as f:
                            f.write(up_tabela.getbuffer())
                        st.success(f"✅ Tabela {fonte_ref} {mes_ano} vinculada a esta obra!")

            # --- ABA 2: MONTAGEM DA PLANILHA ---
            with tab_planilha:
                st.subheader("Planilha Orçamentária Dinâmica")
                st.caption("Adicione, remova ou edite as linhas diretamente na tabela abaixo.")

                # Inicializa os itens se não existirem
                if 'df_orc_dados' not in st.session_state:
                    st.session_state['df_orc_dados'] = pd.DataFrame([
                        {"Item": "1.1", "Descrição": "Serviços Preliminares", "Unidade": "un", "Quantidade": 1, "V. Unitário (R$)": 0.0},
                        {"Item": "2.1", "Descrição": "Infraestrutura", "Unidade": "m3", "Quantidade": 0, "V. Unitário (R$)": 0.0}
                    ])

                # Editor de Dados Interativo
                df_editado = st.data_editor(
                    st.session_state['df_orc_dados'], 
                    num_rows="dynamic", 
                    use_container_width=True,
                    key="editor_orc"
                )

                if st.button("💾 Consolidar e Salvar Planilha"):
                    # Cálculo automático de totais
                    df_editado['V. Total (R$)'] = df_editado['Quantidade'] * df_editado['V. Unitário (R$)']
                    st.session_state['df_orc_dados'] = df_editado
                    
                    # Salva no arquivo Excel
                    path_salvar_orc = os.path.join(pasta_ativa, "03_Orcamento", "ORCAMENTO_CONSOLIDADO.xlsx")
                    df_editado.to_excel(path_salvar_orc, index=False)
                    
                    valor_total = df_editado['V. Total (R$)'].sum()
                    st.success(f"✅ Planilha salva! Valor Total sem BDI: R$ {valor_total:,.2f}")
                    st.info(f"Valor Total com BDI ({bdi_obra}%): R$ {valor_total * (1 + bdi_obra/100):,.2f}")

            # --- ABA 3: CRONOGRAMA FÍSICO-FINANCEIRO DETALHADO ---
            with tab_cron:
                st.subheader("Cronograma Físico-Financeiro Mensal")
                st.write("Distribua a execução de cada item ao longo dos meses.")
                
                if 'df_orc_dados' not in st.session_state or st.session_state['df_orc_dados'].empty:
                    st.warning("⚠️ Primeiro, monte e salve a 'Planilha Orçamentária' na aba ao lado.")
                else:
                    meses_obra = st.number_input("Prazo de Execução (Meses):", 1, 60, 12)
                    
                    # Prepara os dados base (Itens da Planilha)
                    df_base = st.session_state['df_orc_dados'].copy()
                    df_base['V. Total (R$)'] = df_base['Quantidade'] * df_base['V. Unitário (R$)']
                    
                    # Cria as colunas mensais para visualização/edição
                    for m in range(1, meses_obra + 1):
                        df_base[f"Mês {m} (%)"] = 0.0
                    
                    st.write("📝 **Editor de Cronograma (Distribuição Percentual):**")
                    # O usuário preenche o % de execução de cada item em cada mês
                    df_cron_editado = st.data_editor(
                        df_base, 
                        use_container_width=True, 
                        key="editor_cron_detalhado"
                    )
                    
                    if st.button("📊 Consolidar e Gerar Excel do Cronograma"):
                        # Criar o DataFrame final para o Excel com colunas de % e R$
                        lista_final = []
                        for _, linha in df_cron_editado.iterrows():
                            dados_linha = {
                                "Item": linha["Item"],
                                "Descrição": linha["Descrição"],
                                "Unidade": linha["Unidade"],
                                "Quantidade": linha["Quantidade"],
                                "V. Unitário (R$)": linha["V. Unitário (R$)"],
                                "V. Total (R$)": linha["V. Total (R$)"]
                            }
                            # Adiciona colunas intercaladas de % e Valor no Excel
                            for m in range(1, meses_obra + 1):
                                perc = linha[f"Mês {m} (%)"]
                                valor_mes = (perc / 100) * linha["V. Total (R$)"]
                                dados_linha[f"Mês {m} (%)"] = perc
                                dados_linha[f"Mês {m} (R$)"] = valor_mes
                            
                            lista_final.append(dados_linha)
                        
                        df_excel_cron = pd.DataFrame(lista_final)
                        
                        # Salva o arquivo físico
                        path_cron_final = os.path.join(pasta_ativa, "03_Orcamento", "CRONOGRAMA_DETALHADO.xlsx")
                        df_excel_cron.to_excel(path_cron_final, index=False)
                        
                        st.success("✅ Cronograma Detalhado gerado com sucesso!")
                        
                        with open(path_cron_final, "rb") as f:
                            st.download_button(
                                label="📥 Baixar Cronograma em Excel",
                                data=f,
                                file_name=f"Cronograma_{nome_obra}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )

    elif menu == "Acompanhamento":
        import qrcode
        from io import BytesIO

        if 'pasta_ativa' not in st.session_state:
            st.warning("⚠️ Selecione uma obra na barra lateral para acessar a fiscalização.")
        else:
            pasta_ativa = st.session_state['pasta_ativa']
            pasta_fisc = os.path.join(pasta_ativa, "04_Fiscalizacao")
            nome_obra = os.path.basename(pasta_ativa)
            
            st.header(f"📈 Fiscalização e Execução: {nome_obra}")

            # --- O PULO DO GATO: QR CODE NA BARRA LATERAL ---
            with st.sidebar:
                st.divider()
                st.subheader("📱 Transparência QR Code")
                # Link simbólico que aponta para o portal de transparência da prefeitura
                link_pub = f"https://transparencia.gov.br/obras/{nome_obra.replace(' ', '_')}"
                
                qr = qrcode.QRCode(version=1, box_size=10, border=5)
                qr.add_data(link_pub)
                qr.make(fit=True)
                img_qr = qr.make_image(fill_color="black", back_color="white")
                
                buf = BytesIO()
                img_qr.save(buf, format="PNG")
                st.image(buf.getvalue(), caption="QR Code da Placa da Obra")
                st.download_button("📥 Baixar QR Code (PNG)", buf.getvalue(), f"QR_OBRA_{nome_obra}.png")

            # --- ABAS DE TRABALHO ---
            tab_med, tab_diario, tab_fotos = st.tabs([
                "📏 Medições e Aferição", 
                "📝 Diário de Obra Digital", 
                "📸 Relatório Fotográfico"
            ])

            # --- ABA 1: MEDIÇÕES ---
            with tab_med:
                st.subheader("Gestão de Medições")
                st.info("Anexe as planilhas de medição para conferência do fiscal.")
                
                with st.expander("➕ Registrar Medição Mensal"):
                    num_m = st.number_input("Medição nº:", min_value=1, step=1)
                    up_m = st.file_uploader("Upload da Planilha de Medição (Excel):", type=["xlsx"], key="up_med_fisc")
                    status_m = st.selectbox("Parecer do Fiscal:", ["Em Análise", "Aprovada", "Glosa Parcial", "Reprovada"])
                    if st.button("💾 Salvar Registro"):
                        if up_m:
                            caminho_m = os.path.join(pasta_fisc, f"MEDICAO_{num_m}_STATUS_{status_m}.xlsx")
                            with open(caminho_m, "wb") as f:
                                f.write(up_m.getbuffer())
                            st.success(f"Medição {num_m} registrada com sucesso!")
                
                st.divider()
                st.subheader("Histórico de Documentos de Medição")
                arqs_fisc = os.listdir(pasta_fisc)
                meds = [f for f in arqs_fisc if f.startswith("MEDICAO_")]
                if meds:
                    for m in meds:
                        st.write(f"📂 {m}")
                else:
                    st.caption("Nenhum arquivo de medição encontrado.")

            # --- ABA 2: DIÁRIO DE OBRA ---
            with tab_diario:
                st.subheader("Diário de Obra (Lei 14.133)")
                caminho_txt = os.path.join(pasta_fisc, "DIARIO_OBRA.txt")
                
                if os.path.exists(caminho_txt):
                    with open(caminho_txt, "r", encoding="utf-8") as f:
                        st.text_area("Histórico de Ocorrências:", f.read(), height=200, disabled=True)
                
                st.write("---")
                st.write("🖋️ **Novo Registro Diário**")
                col_c1, col_c2, col_c3 = st.columns(3)
                with col_c1: d_ref = st.date_input("Data:")
                with col_c2: clima = st.selectbox("Clima:", ["Ensolarado", "Chuvoso", "Instável"])
                with col_c3: efetivo = st.number_input("Efetivo:", min_value=0)
                
                relato = st.text_area("Descrição dos serviços e intercorrências:")
                if st.button("📝 Assinar Relato"):
                    with open(caminho_txt, "a", encoding="utf-8") as f:
                        f.write(f"\n[{d_ref}] CLIMA: {clima} | EFETIVO: {efetivo}\nRELATO: {relato}\n")
                    st.success("Relato assinado digitalmente e salvo!")
                    st.rerun()

            # --- ABA 3: FOTOS ---
            with tab_fotos:
                st.subheader("Galeria de Evolução da Obra")
                f_up = st.file_uploader("Adicionar fotos (JPG/PNG):", type=["jpg", "png", "jpeg"], accept_multiple_files=True)
                if st.button("🖼️ Salvar Fotos"):
                    if f_up:
                        for f in f_up:
                            with open(os.path.join(pasta_fisc, f.name), "wb") as file:
                                file.write(f.getbuffer())
                        st.rerun()

                st.divider()
                arquivos = os.listdir(pasta_fisc)
                fotos = [f for f in arquivos if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
                if fotos:
                    cols = st.columns(3)
                    for i, foto in enumerate(fotos):
                        with cols[i % 3]:
                            st.image(os.path.join(pasta_fisc, foto), caption=foto)
                else:
                    st.info("Aguardando primeiras fotos da execução.")

if __name__ == "__main__":
    main()